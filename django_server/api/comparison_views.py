from rest_framework.views import APIView
from rest_framework.response import Response
from .models import SecurityLog
from django.http import HttpResponse
from datetime import datetime
from rest_framework.permissions import AllowAny
import csv
import io
import json

def get_comparison_data(file, min_length=2, unique_only=False):
    """
    Helper to parse CSV and find matches in SecurityLog.
    Returns (matches, stats, csv_emails_count, error_msg)
    """
    try:
        # Parse CSV
        decoded_file = file.read().decode('utf-8')
        io_string = io.StringIO(decoded_file)
        reader = csv.DictReader(io_string)
        
        if not reader.fieldnames:
            return None, None, 0, 'Empty CSV file'

        # Robust header matching for email
        fieldnames = reader.fieldnames
        email_key = None
        potential_keys = ['email', 'e-mail', 'recipient', 'address', 'target', 'user']
        
        for key in potential_keys:
            found = next((f for f in fieldnames if key in f.lower()), None)
            if found:
                email_key = found
                break
        
        if not email_key:
            return None, None, 0, 'CSV must contain an "email" column (or similar header like "recipient" or "address")'
        
        csv_emails = set()
        csv_sent_dates = {}
        
        for row in reader:
            val = row.get(email_key)
            if val:
                email = val.strip().lower()
                csv_emails.add(email)
                # Try to capture sent date
                for key in row.keys():
                    if any(date_hint in key.lower() for date_hint in ['sent at', 'sent_at', 'date', 'time']):
                        csv_sent_dates[email] = row[key]
                        break

        # Get security logs
        sec_logs = SecurityLog.objects.all().order_by('-created_at')
        
        matches = []
        unique_matched_emails = set()
        
        for slog in sec_logs:
            if not slog.email:
                continue
                
            email_lower = slog.email.lower()
            
            if email_lower in csv_emails:
                # If unique_only is requested, skip if we already found a match for this email
                if unique_only and email_lower in unique_matched_emails:
                    continue

                # Filter: input_details length >= min_length
                details = slog.input_details or ''
                if len(details) >= min_length:
                    # Optional cleaning for input_details
                    # If it's a JSON string, try to make it prettier
                    cleaned_details = details
                    if details.startswith('{') and details.endswith('}'):
                        try:
                            data = json.loads(details)
                            cleaned_details = ", ".join([f"{k}: {v}" for k, v in data.items()])
                        except:
                            pass
                    
                    matches.append({
                        'id': slog.id,
                        'email': slog.email,
                        'sent_at': csv_sent_dates.get(email_lower),
                        'security_date': slog.created_at,
                        'input_details': details,
                        'cleaned_details': cleaned_details
                    })
                    unique_matched_emails.add(email_lower)

        stats = {
            'total_sent_unique': len(csv_emails),
            'unique_matches': len(unique_matched_emails),
            'total_matches': len(matches)
        }
        
        return matches, stats, len(csv_emails), None
    except Exception as e:
        return None, None, 0, str(e)

class ComparisonAnalyticsView(APIView):
    def post(self, request):
        file = request.FILES.get('file')
        if not file:
            return Response({'error': 'No file uploaded'}, status=400)

        min_length = int(request.data.get('min_length', 2))
        unique_only = request.data.get('unique_only', 'false').lower() == 'true'
        
        matches, stats, _, error = get_comparison_data(file, min_length, unique_only)
        if error:
            return Response({'error': error}, status=400 if 'column' in error else 500)

        return Response({
            'matches': matches,
            'stats': stats
        })

class ComparisonExportView(APIView):
    """Word Document Export"""
    authentication_classes = []
    permission_classes = [AllowAny]
    
    def post(self, request):
        try:
            file = request.FILES.get('file')
            if not file:
                return Response({'error': 'No file uploaded'}, status=400)

            min_length = int(request.data.get('min_length', 2))
            unique_only = request.data.get('unique_only', 'false').lower() == 'true'
            matches, stats, total_emails, error = get_comparison_data(file, min_length, unique_only)
            
            if error:
                return Response({'error': error}, status=400 if 'column' in error else 500)

            # Word Document Generation
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            
            doc = Document()
            
            # Helper to add shading
            def set_cell_shading(cell, color):
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
                cell._element.get_or_add_tcPr().append(shading_elm)

            # Title & Header
            title = doc.add_heading('Deliverability Diagnostic Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.color.rgb = RGBColor(79, 70, 229) # Indigo
            
            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta.add_run(f"System Reconciliation: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
            
            # Executive Summary Section
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph('This report provides a cross-reconciliation between external campaign reach and internal security log identifiers.')
            
            # Stats Grid / KPI section
            stats_table = doc.add_table(rows=2, cols=2)
            stats_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            confidence = (len(matches) / total_emails * 100) if total_emails > 0 else 0
            
            kpis = [
                ('Campaign size', f"{total_emails} Unique Recipients"),
                ('Security Hits', f"{len(matches)} Identified Events"),
                ('Confidence Score', f"{confidence:.1f}% Match Velocity"),
                ('Data Coverage', f"{stats['unique_matches']} Unique Identifiers")
            ]
            
            cells = stats_table._cells
            for i, (label, val) in enumerate(kpis):
                cell = cells[i]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(f"{label.upper()}\n").font.size = Pt(9)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(100, 116, 139) # Slate
                p.add_run(val).font.size = Pt(14)
                p.runs[1].font.bold = True
                p.runs[1].font.color.rgb = RGBColor(79, 70, 229) # Indigo
            
            doc.add_paragraph() # Spacing
            
            # Diagnostic Audit Log
            doc.add_heading('Diagnostic Audit Log', level=1)
            doc.add_paragraph(f"Detailed list of records identified in system logs with input details >= {min_length} characters.")
            
            # Match Table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            headers = ['SL NO', 'TARGET IDENTITY', 'SENT DATE', 'LOG TIMESTAMP', 'SYSTEM DETAILS']
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                cell = hdr_cells[i]
                p = cell.paragraphs[0]
                p.text = h
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                set_cell_shading(cell, "4F46E5") # Indigo primary

            for idx, m in enumerate(matches):
                row_cells = table.add_row().cells
                row_cells[0].text = f"{(idx + 1):02d}"
                row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                
                row_cells[1].text = m['email']
                row_cells[1].paragraphs[0].runs[0].font.bold = True
                
                row_cells[2].text = str(m['sent_at']) if m['sent_at'] else '-'
                row_cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                
                row_cells[3].text = m['security_date'].strftime('%Y-%m-%d %H:%M:%S')
                row_cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                
                details_p = row_cells[4].paragraphs[0]
                details_p.text = m['input_details'] or '-'
                details_p.runs[0].font.size = Pt(8)
                details_p.runs[0].font.name = 'Courier New'

            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph("End of Deliverability Diagnostic Report")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.runs[0].italic = True
            
            # Save
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="deliverability_report_{datetime.now().strftime("%Y%m%d_%H%M")}.docx"'
            return response
            
        except Exception as e:
            return Response({'error': str(e)}, status=500)

class ComparisonCSVExportView(APIView):
    """CSV Export"""
    authentication_classes = []
    permission_classes = [AllowAny]
    
    def post(self, request):
        try:
            file = request.FILES.get('file')
            if not file:
                return Response({'error': 'No file uploaded'}, status=400)

            min_length = int(request.data.get('min_length', 2))
            unique_only = request.data.get('unique_only', 'false').lower() == 'true'
            matches, stats, total_emails, error = get_comparison_data(file, min_length, unique_only)
            
            if error:
                return Response({'error': error}, status=400 if 'column' in error else 500)

            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = 'attachment; filename="comparison_analysis.csv"'
            
            writer = csv.writer(response)
            writer.writerow(['Email Address', 'Email Sent Date', 'Security Log Date', 'Input Details'])
            
            for m in matches:
                writer.writerow([
                    m['email'],
                    m['sent_at'] or '-',
                    m['security_date'].strftime('%Y-%m-%d %H:%M:%S'),
                    m['input_details'] or '-'
                ])
                
            return response
        except Exception as e:
            return Response({'error': str(e)}, status=500)
