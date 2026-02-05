import csv
import io
import json
import os
from django.conf import settings
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.db import transaction
from django.db.models import Sum, Count, Q, F
from django.utils import timezone
from django.utils.dateparse import parse_datetime
from .models import MailConfig, EmailRecipient, EmailTemplate, EmailCampaign, EmailLog, SecurityLog
from .serializers import MailConfigSerializer, EmailRecipientSerializer, EmailTemplateSerializer, EmailCampaignSerializer, EmailLogSerializer
from .utils.mailer import process_campaign, send_email, process_template_variables
import uuid
from django.http import HttpResponse, HttpResponseRedirect
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import logging

# Configure Logger for API views to write to mailer.log
logger = logging.getLogger('api_views')
if not logger.handlers:
    log_path = os.path.join(settings.BASE_DIR, 'mailer.log')
    handler = logging.FileHandler(log_path)
    formatter = logging.Formatter('%(asctime)s - [API] - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    logger.setLevel(logging.INFO)

# --- Config ---
class ConfigView(APIView):
    def get(self, request):
        config = MailConfig.objects.first()
        if config:
            serializer = MailConfigSerializer(config)
            return Response(serializer.data)
        return Response({})

    def post(self, request):
        config = MailConfig.objects.first()
        data = request.data
        
        try:
            # If updating, merge with existing instance
            if config:
                serializer = MailConfigSerializer(config, data=data, partial=True)
            else:
                data['id'] = str(uuid.uuid4())
                serializer = MailConfigSerializer(data=data)

            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# --- Recipients ---
class RecipientListView(APIView):
    def get(self, request):
        recipients = EmailRecipient.objects.all().order_by('-created_at')
        serializer = EmailRecipientSerializer(recipients, many=True)
        return Response(serializer.data)

    def post(self, request):
        data = request.data
        if 'id' not in data:
            data['id'] = str(uuid.uuid4())
            
        serializer = EmailRecipientSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class RecipientDetailView(APIView):
    def delete(self, request, pk):
        try:
            recipient = EmailRecipient.objects.get(pk=pk)
            recipient.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except EmailRecipient.DoesNotExist:
            return Response({'error': 'Recipient not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class RecipientBulkView(APIView):
    def post(self, request):
        recipients_data = request.data
        if not isinstance(recipients_data, list) or len(recipients_data) == 0:
            return Response({'error': 'Invalid input: expected non-empty array of recipients'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Optimized approach using bulk_create
            recipient_objects = []
            for r in recipients_data:
                if not r.get('email'):
                    continue
                recipient_objects.append(EmailRecipient(
                    id=str(uuid.uuid4()),
                    email=r['email'],
                    name=r.get('name')
                ))
            
            # bulk_create with ignore_conflicts=True handles "ON CONFLICT DO NOTHING" (Postgres/SQLite)
            EmailRecipient.objects.bulk_create(recipient_objects, ignore_conflicts=True)
            
            return Response({'success': True, 'count': len(recipients_data)})
        except Exception as e:
            return Response({'error': 'Failed to bulk add recipients', 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# --- Templates ---
class TemplateListView(APIView):
    def get(self, request):
        templates = EmailTemplate.objects.all().order_by('-created_at')
        serializer = EmailTemplateSerializer(templates, many=True)
        return Response(serializer.data)

    def post(self, request):
        data = request.data
        if 'id' not in data:
            data['id'] = str(uuid.uuid4())
        serializer = EmailTemplateSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class TemplateDetailView(APIView):
    def delete(self, request, pk):
        try:
            template = EmailTemplate.objects.get(pk=pk)
            template.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except EmailTemplate.DoesNotExist:
            return Response({'error': 'Template not found'}, status=status.HTTP_404_NOT_FOUND)

# --- Campaigns ---
class CampaignListView(APIView):
    def get(self, request):
        campaigns = EmailCampaign.objects.all().order_by('-created_at')
        serializer = EmailCampaignSerializer(campaigns, many=True)
        return Response(serializer.data)

    def post(self, request):
        data = request.data
        if 'id' not in data:
            data['id'] = str(uuid.uuid4())
        
        # If scheduled_at is present, set status to scheduled
        if data.get('scheduled_at'):
            data['status'] = 'scheduled'
        else:
            data['status'] = 'draft'
            
        data['total_recipients'] = len(data.get('recipient_ids', []))
        
        serializer = EmailCampaignSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class CampaignDetailView(APIView):
    def get(self, request, pk):
        try:
            campaign = EmailCampaign.objects.get(pk=pk)
            serializer = EmailCampaignSerializer(campaign)
            return Response(serializer.data)
        except EmailCampaign.DoesNotExist:
            return Response({'error': 'Campaign not found'}, status=status.HTTP_404_NOT_FOUND)

    def delete(self, request, pk):
        try:
            campaign = EmailCampaign.objects.get(pk=pk)
            logger.info(f"API: Deleting campaign {pk} ({campaign.name}). Any running processes will abort.")
            campaign.delete()
            logger.info(f"API: Campaign {pk} deleted successfully")
            return Response(status=status.HTTP_204_NO_CONTENT)
        except EmailCampaign.DoesNotExist:
            return Response({'error': 'Campaign not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            logger.error(f"API: Error deleting campaign {pk}: {e}")
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class CampaignStatusView(APIView):
    def patch(self, request, pk):
        try:
            campaign = EmailCampaign.objects.get(pk=pk)
            status_val = request.data.get('status')
            if status_val:
                campaign.status = status_val
                campaign.save()
                return Response({'success': True, 'status': status_val})
            return Response({'error': 'Status required'}, status=status.HTTP_400_BAD_REQUEST)
        except EmailCampaign.DoesNotExist:
            return Response({'error': 'Campaign not found'}, status=status.HTTP_404_NOT_FOUND)

class CampaignStartView(APIView):
    def post(self, request, pk):
        try:
            logger.info(f"API: Received start request for campaign {pk}")
            
            # Get the campaign
            campaign = EmailCampaign.objects.get(pk=pk)
            
            # Get mode from request (default to 'start' for immediate sending)
            mode = request.data.get('mode', 'start')
            
            if mode not in ['draft', 'start', 'pause']:
                return Response({'error': 'Invalid mode. Must be draft, start, or pause'}, 
                              status=status.HTTP_400_BAD_REQUEST)
            
            logger.info(f"API: Campaign {pk} mode: {mode}")
            
            if mode == 'draft':
                # Just save as draft, don't start sending
                campaign.status = 'draft'
                campaign.save()
                logger.info(f"API: Campaign {pk} saved as draft")
                return Response({'success': True, 'message': 'Campaign saved as draft', 'status': 'draft'})
            
            elif mode == 'pause':
                # Set to paused state
                campaign.status = 'paused'
                campaign.save()
                logger.info(f"API: Campaign {pk} paused")
                return Response({'success': True, 'message': 'Campaign paused', 'status': 'paused'})
            
            else:  # mode == 'start'
                # Start sending immediately
                base_url = f"{request.scheme}://{request.get_host()}"
                # Start in background, but only after transaction commits to ensure thread sees data
                transaction.on_commit(lambda: process_campaign(pk, base_url))
                logger.info(f"API: Campaign {pk} queued for background processing")
                return Response({'success': True, 'message': 'Campaign started', 'status': 'sending'})
                
        except EmailCampaign.DoesNotExist:
            return Response({'error': 'Campaign not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            logger.error(f"API: Failed to start campaign {pk}: {e}")
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# --- Logs ---
class CampaignLogsView(APIView):
    def get(self, request, pk):
        logs = EmailLog.objects.filter(campaign_id=pk).order_by('-created_at')
        serializer = EmailLogSerializer(logs, many=True)
        return Response(serializer.data)

class CampaignServerLogsView(APIView):
    def get(self, request, pk):
        try:
            log_path = os.path.join(settings.BASE_DIR, 'mailer.log')
            scheduler_log_path = os.path.join(settings.BASE_DIR, 'scheduler.log')
            
            logs = []
            
            # Helper to read logs
            def read_logs(path, source):
                try:
                    if not os.path.exists(path):
                         # If file doesn't exist, just skip
                        return
                    
                    # Check file permissions simply by trying to open
                    # Check file size, if too big, read last 1MB
                    file_size = os.path.getsize(path)
                    read_mode = 'r'
                    
                    with open(path, read_mode, encoding='utf-8', errors='ignore') as f:
                        if file_size > 512 * 1024: # Limit to last 512KB for speed
                            f.seek(file_size - 512 * 1024)
                        
                        for line in f:
                            # Filter for logs containing the campaign ID
                            if pk in line:
                                logs.append({
                                    'source': source,
                                    'message': line.strip(),
                                    'timestamp': line.split(' - ')[0] if ' - ' in line else ''
                                })
                except PermissionError:
                    logs.append({
                        'source': source,
                        'message': f"ERROR: Permission denied reading {os.path.basename(path)}. Please fix server permissions.",
                        'timestamp': timezone.now().isoformat()
                    })
                except Exception as e:
                    logs.append({
                        'source': source,
                        'message': f"ERROR: Failed to read log: {str(e)}",
                        'timestamp': timezone.now().isoformat()
                    })

            read_logs(log_path, 'mailer')
            read_logs(scheduler_log_path, 'scheduler')
            
            # Sort by timestamp (approximate) or just return as is (lines are chronological per file)
            # Merging two files chronologically is hard without parsing dates strictly.
            # Let's just return them.
            
            return Response(logs)
        except Exception as e:
            return Response({'error': str(e)}, status=500)

class ClearLogsView(APIView):
    """Clear server log files (mailer.log and scheduler.log)"""
    def post(self, request):
        try:
            log_path = os.path.join(settings.BASE_DIR, 'mailer.log')
            scheduler_log_path = os.path.join(settings.BASE_DIR, 'scheduler.log')
            
            cleared = []
            errors = []
            
            for path, name in [(log_path, 'mailer.log'), (scheduler_log_path, 'scheduler.log')]:
                try:
                    if os.path.exists(path):
                        # Clear file by opening in write mode
                        with open(path, 'w') as f:
                            f.write('')
                        cleared.append(name)
                        logger.info(f"API: Cleared log file: {name}")
                    else:
                        errors.append(f"{name} does not exist")
                except PermissionError:
                    errors.append(f"Permission denied for {name}")
                except Exception as e:
                    errors.append(f"Failed to clear {name}: {str(e)}")
            
            if cleared and not errors:
                return Response({
                    'success': True,
                    'message': f'Cleared {", ".join(cleared)}',
                    'cleared': cleared
                })
            elif cleared and errors:
                return Response({
                    'success': True,
                    'message': f'Partially cleared. Errors: {", ".join(errors)}',
                    'cleared': cleared,
                    'errors': errors
                })
            else:
                return Response({
                    'success': False,
                    'message': 'Failed to clear logs',
                    'errors': errors
                }, status=500)
                
        except Exception as e:
            logger.error(f"API: Error clearing logs: {e}")
            return Response({'error': str(e)}, status=500)


# --- Stats ---
class StatsSummaryView(APIView):
    def get(self, request):
        campaign_id = request.GET.get('campaignId')
        
        # Base querysets
        campaigns_qs = EmailCampaign.objects.all()
        logs_qs = EmailLog.objects.all()
        recipients_qs = EmailRecipient.objects.all()

        if campaign_id and campaign_id != 'all':
            campaigns_qs = campaigns_qs.filter(id=campaign_id)
            logs_qs = logs_qs.filter(campaign_id=campaign_id)
            # For a single campaign, total recipients is the number of logs (or recipients in campaign)
            # Usually users expect "Recipients targeted" for a campaign stats view.
            # But let's stick to aggregates.

        # Aggregates
        total_predictions = campaigns_qs.aggregate(
            sent=Sum('sent_count'),
            failed=Sum('failed_count'),
            opens=Sum('open_count'),
            clicks=Sum('click_count')
        )
        
        total_sent = total_predictions['sent'] or 0
        total_failed = total_predictions['failed'] or 0
        total_opens = total_predictions['opens'] or 0
        total_clicks = total_predictions['clicks'] or 0
        
        # Unique counts (from logs)
        unique_opens = logs_qs.exclude(opened_at__isnull=True).values('recipient').distinct().count()
        unique_clicks = logs_qs.exclude(clicked_at__isnull=True).values('recipient').distinct().count()
        
        total_recipients = recipients_qs.count() # This is global total recipients in DB
        if campaign_id and campaign_id != 'all':
             # If filtering by campaign, total recipients = pending + sent + failed for that campaign (approximation)
             total_recipients = logs_qs.count() 

        total_pending = logs_qs.filter(status='pending').count()
        
        # Daily Stats (last 7 days)
        from django.db.models.functions import TruncDate
        from datetime import timedelta
        
        seven_days_ago = timezone.now() - timedelta(days=7)
        daily_stats_qs = logs_qs.filter(created_at__gte=seven_days_ago).annotate(date=TruncDate('created_at')).values('date').annotate(
            sent=Count('id', filter=Q(status='sent')),
            failed=Count('id', filter=Q(status='failed')),
            opens=Count('id', filter=Q(opened_at__isnull=False)),
            clicks=Count('id', filter=Q(clicked_at__isnull=False))
        ).order_by('date')
        
        daily_stats = []
        for stat in daily_stats_qs:
            daily_stats.append({
                'date': stat['date'].isoformat(),
                'sent': stat['sent'],
                'failed': stat['failed'],
                'opens': stat['opens'],
                'clicks': stat['clicks']
            })

        return Response({
            'totalSent': total_sent,
            'totalFailed': total_failed,
            'totalPending': total_pending,
            'totalRecipients': total_recipients,
            'totalOpens': total_opens,
            'uniqueOpens': unique_opens,
            'totalClicks': total_clicks,
            'uniqueClicks': unique_clicks,
            'totalCampaigns': EmailCampaign.objects.count(), # Always total campaigns count
            'dailyStats': daily_stats
        })

# --- Test Email ---
class TestEmailView(APIView):
    def post(self, request):
        to = request.data.get('to')
        html = request.data.get('html')
        subject = request.data.get('subject')
        
        config = MailConfig.objects.first()
        if not config or not config.is_configured:
            return Response({'error': 'Email provider not configured'}, status=status.HTTP_400_BAD_REQUEST)
        
        recipient_name = 'Test User'
        # Pass empty log_id for test
        processed_html = process_template_variables(html, to, recipient_name, 'test-id', False, '')
        
        result = send_email(to, subject, processed_html, config)
        
        if result['success']:
            return Response({'success': True})
        else:
            print(f"Test Email Error: {result.get('error')}")
            return Response({'error': result.get('error')}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# --- Tracking Pixel ---
from django.http import HttpResponse
import base64

class TrackEmailView(APIView):
    def get(self, request):
        log_id = request.GET.get('id')
        type_ = request.GET.get('type')
        url = request.GET.get('url')  # For click tracking
        
        if log_id:
            try:
                log = EmailLog.objects.get(id=log_id)
                campaign = log.campaign
                
                # Capture user agent and IP
                user_agent = request.META.get('HTTP_USER_AGENT', '')
                ip_address = request.META.get('HTTP_X_FORWARDED_FOR', request.META.get('REMOTE_ADDR', ''))
                if ',' in ip_address:
                    ip_address = ip_address.split(',')[0].strip()
                
                if type_ == 'open' and not log.opened_at:
                    log.opened_at = timezone.now()
                    log.user_agent = user_agent[:500]  # Limit to field size
                    log.ip_address = ip_address
                    log.save()
                    # Atomic increment
                    EmailCampaign.objects.filter(pk=campaign.pk).update(open_count=F('open_count') + 1)
                    logger.info(f"Email opened: {log.recipient_email} (Campaign: {campaign.name})")
                
                elif type_ == 'click' and not log.clicked_at:
                    log.clicked_at = timezone.now()
                    if not log.user_agent:  # Only set if not already set by open
                        log.user_agent = user_agent[:500]
                        log.ip_address = ip_address
                    log.save()
                    # Atomic increment
                    EmailCampaign.objects.filter(pk=campaign.pk).update(click_count=F('click_count') + 1)
                    logger.info(f"Email clicked: {log.recipient_email} (Campaign: {campaign.name})")
                    
                    # Redirect to actual URL if provided
                    if url:
                        return HttpResponseRedirect(url)
                    
            except EmailLog.DoesNotExist:
                logger.warning(f"Tracking attempt for non-existent log: {log_id}")
            except Exception as e:
                logger.error(f"Tracking error: {e}")
                pass

        # Return 1x1 transparent GIF for open tracking
        pixel = base64.b64decode('R0lGODlhAQABAIAAAAAAAP///yH5BAEAAAAALAAAAAABAAEAAAIBRAA7')
        return HttpResponse(pixel, content_type='image/gif')

# --- Recent Activity ---
class RecentActivityView(APIView):
    def get(self, request):
        # Fetch last 20 logs
        logs = EmailLog.objects.filter(
            status__in=['sent', 'failed']
        ).select_related('campaign').order_by('-created_at')[:20]
        
        data = []
        for log in logs:
            data.append({
                'id': log.id,
                'email': log.recipient,
                'campaign_name': log.campaign.name if log.campaign else 'Test Email',
                'status': log.status,
                'created_at': log.created_at
            })
            
        return Response(data)

# --- Security Analytics ---
class SecurityLogImportView(APIView):
    def post(self, request):
        if 'file' not in request.FILES:
            return Response({'error': 'No file provided'}, status=400)
        
        file = request.FILES['file']
        if not file.name.endswith('.csv'):
            return Response({'error': 'File must be CSV'}, status=400)
            
        decoded_file = file.read().decode('utf-8-sig') # Handle BOM
        io_string = io.StringIO(decoded_file)
        # Use a DictReader that handles potential whitespace in headers
        reader = csv.DictReader(io_string)
        
        # Normalize headers: strip whitespace
        if reader.fieldnames:
            reader.fieldnames = [name.strip() for name in reader.fieldnames]

        try:
            logs = []
            for row in reader:
                email = row.get('email')
                # Fallback for email if empty (some logs might be IP only)
                if not email:
                    continue
                    
                created_at_str = row.get('created_at')
                created_at = timezone.now()
                if created_at_str:
                    parsed_date = parse_datetime(created_at_str)
                    if parsed_date:
                        created_at = parsed_date

                logs.append(SecurityLog(
                    email=email,
                    ip_address=row.get('ip_address'),
                    user_agent=row.get('user_agent'),
                    input_details=row.get('input_details'),
                    attempt_status=row.get('status'), # Map CSV 'status' to attempt_status
                    created_at=created_at
                ))
                
            if logs:
                SecurityLog.objects.bulk_create(logs)
                
            return Response({'message': f'Imported {len(logs)} logs successfully'})
        except Exception as e:
            return Response({'error': f'Import failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class SecurityLogBulkView(APIView):
    def post(self, request):
        try:
            action = request.data.get('action')
            print(f"DEBUG: Bulk action: {action}, Data: {request.data}")
            
            if action == 'delete_all':
                count, _ = SecurityLog.objects.all().delete()
                return Response({'message': f'Deleted {count} logs'})
                
            elif action == 'delete_selected':
                ids = request.data.get('ids', [])
                count, _ = SecurityLog.objects.filter(id__in=ids).delete()
                return Response({'message': f'Deleted {count} logs'})
                
            elif action == 'update_status':
                ids = request.data.get('ids', [])
                new_status = request.data.get('status')
                if not new_status:
                    return Response({'error': 'Status required'}, status=400)
                
                print(f"DEBUG: Updating {len(ids)} logs to {new_status}")
                count = SecurityLog.objects.filter(id__in=ids).update(review_status=new_status)
                return Response({'message': f'Updated {count} logs'})
                
            return Response({'error': 'Invalid action'}, status=400)
            
        except Exception as e:
            print(f"ERROR: Bulk action failed: {str(e)}")
            return Response({'error': str(e)}, status=500)

class SecurityLogAnalyticsView(APIView):
    def get(self, request):
        try:
            # Top IPs
            top_ips = list(SecurityLog.objects.values('ip_address')
                .annotate(count=Count('ip_address'))
                .order_by('-count')[:5])
            
            # Input Length Distribution
            # We want to highlight inputs > 2 characters
            all_inputs = SecurityLog.objects.all().values_list('input_details', flat=True)
            long_input_count = 0
            short_input_count = 0
            
            for inp in all_inputs:
                if inp and len(str(inp)) > 2:
                    long_input_count += 1
                else:
                    short_input_count += 1
            
            status_dist = [
                {'name': 'Detailed Input (>2 chars)', 'count': long_input_count},
                {'name': 'Minimal Input', 'count': short_input_count}
            ]
                
            # Timeline data for Scaler Graph (Area/Line) - Last 7 Days
            # SQLite doesn't support TruncDate easily in all versions, but basic date grouping works
            # Using raw SQL or Python aggregation might be safer if DB is simple, but let's try ORM first
            # Simplification: Fetch date and count in python to avoid DB complexity issues with sqlite date functions
            from django.utils import timezone
            import datetime
            last_week = timezone.now() - datetime.timedelta(days=7)
            timeline_logs = SecurityLog.objects.filter(created_at__gte=last_week).values('created_at', 'input_details', 'email')
            
            # Aggregate in python
            date_counts = {}
            for log in timeline_logs:
                date_str = log['created_at'].strftime('%Y-%m-%d')
                has_input = bool(log.get('input_details'))
                has_email = bool(log.get('email'))
                
                if date_str not in date_counts:
                    date_counts[date_str] = {'date': date_str, 'email_count': 0, 'input_count': 0}
                
                if has_email:
                    date_counts[date_str]['email_count'] += 1
                
                if has_input:
                    date_counts[date_str]['input_count'] += 1
                    
            timeline_data = sorted(list(date_counts.values()), key=lambda x: x['date'])

            # Filters and Pagination
            page = int(request.query_params.get('page', 1))
            limit = int(request.query_params.get('limit', 1000))
            search_query = request.query_params.get('search', '').lower()
            date_filter = request.query_params.get('date', '')
            
            print(f"DEBUG: Search: '{search_query}', Date: '{date_filter}'")
            
            logs_query = SecurityLog.objects.all().order_by('-created_at')
            
            if search_query:
                logs_query = logs_query.filter(
                    Q(email__icontains=search_query) | 
                    Q(ip_address__icontains=search_query) |
                    Q(input_details__icontains=search_query)
                )
                
            if date_filter:
                try:
                    search_date = timezone.datetime.strptime(date_filter, '%Y-%m-%d').date()
                    logs_query = logs_query.filter(created_at__date=search_date)
                except ValueError:
                    pass

            total_logs = logs_query.count()
            success_count = logs_query.filter(attempt_status='success').count()
            failure_count = logs_query.filter(attempt_status='failure').count()

            start = (page - 1) * limit
            end = start + limit
            recent_logs = logs_query[start:end]
            
            logs_data = [{
                'id': log.id,
                'email': log.email,
                'ip_address': log.ip_address,
                'user_agent': log.user_agent,
                'created_at': log.created_at,
                'input_details': log.input_details,
                'attempt_status': log.attempt_status,
                'review_status': log.review_status
            } for log in recent_logs]
            
            return Response({
                'top_ips': top_ips,
                'status_distribution': status_dist,
                'timeline_data': timeline_data,
                'recent_logs': logs_data,
                'counts': {
                    'total': total_logs,
                    'success': success_count,
                    'failure': failure_count
                },
                'pagination': {
                    'total': total_logs,
                    'page': page,
                    'limit': limit,
                    'pages': (total_logs + limit - 1) // limit
                }
            })
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class SecurityLogActionView(APIView):
    def delete(self, request, pk):
        try:
            log = SecurityLog.objects.get(pk=pk)
            log.delete()
            return Response({'message': 'Log deleted'})
        except SecurityLog.DoesNotExist:
            return Response({'error': 'Log not found'}, status=404)

    def patch(self, request, pk):
        try:
            log = SecurityLog.objects.get(pk=pk)
            status = request.data.get('review_status')
            if status:
                log.review_status = status
                log.save()
                return Response({'message': 'Status updated'})
            return Response({'error': 'No status provided'}, status=400)
        except SecurityLog.DoesNotExist:
            return Response({'error': 'Log not found'}, status=404)

# --- Export Views ---
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from django.http import HttpResponse

from rest_framework.permissions import AllowAny

class StatsExportView(APIView):
    authentication_classes = []
    permission_classes = [AllowAny]
    def get(self, request):
        export_type = request.query_params.get('type', 'csv')
        campaign_id = request.query_params.get('campaign_id')
        
        if campaign_id:
            # Export Combined Logs for specific campaign
            try:
                campaign = EmailCampaign.objects.get(id=campaign_id)
                logs = EmailLog.objects.filter(campaign=campaign).order_by('-created_at')
                
                if export_type == 'csv':
                    response = HttpResponse(content_type='text/csv')
                    filename = f"campaign_logs_{campaign.name.replace(' ', '_')}.csv"
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    
                    writer = csv.writer(response)
                    writer.writerow(['Campaign Name', campaign.name])
                    writer.writerow(['Date', campaign.created_at])
                    writer.writerow([])
                    writer.writerow(['Recipient Email', 'Status', 'Sent At', 'Error Message'])
                    
                    for log in logs:
                        writer.writerow([
                            log.recipient_email,
                            log.status.upper(),
                            log.sent_at if log.sent_at else '-',
                            log.error_message if log.error_message else '-'
                        ])
                    return response
            except EmailCampaign.DoesNotExist:
                return Response({'error': 'Campaign not found'}, status=404)
        
        # Aggregate stats (simplified for report)
        stats = {
            'total_sent': EmailLog.objects.filter(status='sent').count(),
            'total_failed': EmailLog.objects.filter(status='failed').count(),
            'total_campaigns': EmailCampaign.objects.count(),
            'recent_campaigns': EmailCampaign.objects.all().order_by('-created_at')[:10]
        }

        if export_type == 'csv':
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = 'attachment; filename="campaign_stats.csv"'
            
            writer = csv.writer(response)
            writer.writerow(['Metric', 'Value'])
            writer.writerow(['Total Sent', stats['total_sent']])
            writer.writerow(['Total Failed', stats['total_failed']])
            writer.writerow(['Total Campaigns', stats['total_campaigns']])
            writer.writerow([])
            writer.writerow(['Recent Campaigns'])
            writer.writerow(['Name', 'Status', 'Sent', 'Failed', 'Date'])
            for c in stats['recent_campaigns']:
                writer.writerow([c.name, c.status, c.sent_count, c.failed_count, c.created_at])
            
            return response

        elif export_type == 'pdf':
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="campaign_stats.pdf"'
            
            doc = SimpleDocTemplate(response, pagesize=letter)
            elements = []
            styles = getSampleStyleSheet()
            
            elements.append(Paragraph("Campaign Analytics Report", styles['Title']))
            
            # Summary Table
            data = [
                ['Metric', 'Value'],
                ['Total Sent', stats['total_sent']],
                ['Total Failed', stats['total_failed']],
                ['Total Campaigns', stats['total_campaigns']]
            ]
            t = Table(data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            elements.append(t)
            
            doc.build(elements)
            return response
        
        elif export_type == 'docx':
            # Word document export
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime
            
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            if campaign_id:
                # Campaign-specific report
                filename = f"campaign_report_{campaign.name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                
                # Title
                title = doc.add_heading(f'Campaign Report: {campaign.name}', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title.runs[0]
                title_run.font.color.rgb = RGBColor(0, 102, 204)
                
                # Campaign Summary Section
                doc.add_heading('Campaign Summary', level=1)
                summary_table = doc.add_table(rows=5, cols=2)
                summary_table.style = 'Light Grid Accent 1'
                
                summary_data = [
                    ('Campaign Name', campaign.name),
                    ('Subject', campaign.subject),
                    ('Status', campaign.status.upper()),
                    ('Created', campaign.created_at.strftime('%Y-%m-%d %H:%M')),
                    ('Sent', campaign.sent_at.strftime('%Y-%m-%d %H:%M') if campaign.sent_at else 'Not sent yet')
                ]
                
                for i, (label, value) in enumerate(summary_data):
                    row = summary_table.rows[i]
                    row.cells[0].text = label
                    row.cells[1].text = str(value)
                    # Bold labels
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                
                doc.add_paragraph()  # Spacing
                
                # Statistics Section
                doc.add_heading('Statistics', level=1)
                stats_table = doc.add_table(rows=5, cols=2)
                stats_table.style = 'Light Grid Accent 1'
                
                stats_data = [
                    ('Total Recipients', campaign.total_recipients),
                    ('Emails Sent', campaign.sent_count),
                    ('Failed', campaign.failed_count),
                    ('Opened', campaign.open_count),
                    ('Clicked', campaign.click_count)
                ]
                
                for i, (label, value) in enumerate(stats_data):
                    row = stats_table.rows[i]
                    row.cells[0].text = label
                    row.cells[1].text = str(value)
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                    
                    # Color code based on metric
                    if label == 'Failed' and value > 0:
                        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 53, 69)
                    elif label in ['Emails Sent', 'Opened', 'Clicked']:
                        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(40, 167, 69)
                
                doc.add_paragraph()  # Spacing
                
                # Recipient Details Section
                doc.add_heading('Recipient Details', level=1)
                
                # Create recipient table
                recipient_table = doc.add_table(rows=1, cols=4)
                recipient_table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = recipient_table.rows[0].cells
                headers = ['Email', 'Status', 'Sent At', 'Error']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                    header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    # Blue background for header
                    from docx.oxml.ns import nsdecls
                    from docx.oxml import parse_xml
                    shading_elm = parse_xml(r'<w:shd {} w:fill="0066CC"/>'.format(nsdecls('w')))
                    header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                
                # Add recipient rows
                for log in logs:
                    row_cells = recipient_table.add_row().cells
                    row_cells[0].text = log.recipient_email
                    row_cells[1].text = log.status.upper()
                    row_cells[2].text = log.sent_at.strftime('%Y-%m-%d %H:%M') if log.sent_at else '-'
                    row_cells[3].text = log.error_message[:50] if log.error_message else '-'
                    
                    # Color code status
                    if log.status == 'sent':
                        row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(40, 167, 69)
                    elif log.status == 'failed':
                        row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 53, 69)
            
            else:
                # Overall stats report
                filename = f"campaign_stats_{datetime.now().strftime('%Y%m%d')}.docx"
                
                title = doc.add_heading('Campaign Analytics Report', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title.runs[0]
                title_run.font.color.rgb = RGBColor(0, 102, 204)
                
                doc.add_heading('Overall Statistics', level=1)
                stats_table = doc.add_table(rows=3, cols=2)
                stats_table.style = 'Light Grid Accent 1'
                
                overall_stats = [
                    ('Total Sent', stats['total_sent']),
                    ('Total Failed', stats['total_failed']),
                    ('Total Campaigns', stats['total_campaigns'])
                ]
                
                for i, (label, value) in enumerate(overall_stats):
                    row = stats_table.rows[i]
                    row.cells[0].text = label
                    row.cells[1].text = str(value)
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            # Footer
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run(f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Save to response
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
        return Response({'error': 'Invalid type'}, status=400)

class SecurityLogExportView(APIView):
    authentication_classes = []
    permission_classes = [AllowAny]
    def get(self, request):
        export_type = request.query_params.get('type', 'csv')
        status_filter = request.query_params.get('status')
        
        logs = SecurityLog.objects.all().order_by('-created_at')
        if status_filter:
            logs = logs.filter(attempt_status=status_filter)
            
        if export_type == 'csv':
            response = HttpResponse(content_type='text/csv')
            filename = f"security_logs_{status_filter if status_filter else 'all'}.csv"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            writer = csv.writer(response)
            writer.writerow(['Time', 'Email', 'IP Address', 'Attempt Status', 'Review Status', 'User Agent', 'Input Details'])
            
            seen_emails = set()
            for log in logs:
                if log.email not in seen_emails:
                    writer.writerow([
                        log.created_at, log.email, log.ip_address, 
                        log.attempt_status, log.review_status, log.user_agent,
                        log.input_details
                    ])
                    seen_emails.add(log.email)
            return response
            
        elif export_type == 'pdf':
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="security_logs.pdf"'
            
            doc = SimpleDocTemplate(response, pagesize=letter)
            elements = []
            styles = getSampleStyleSheet()
            
            elements.append(Paragraph("Security Logs Report", styles['Title']))
            
            # Table Data with wrapping for long text
            data = [['Time', 'Email', 'IP', 'Status']]
            for log in logs[:50]: # Limit for PDF readability
                data.append([
                    str(log.created_at)[:19],
                    log.email,
                    log.ip_address,
                    log.attempt_status
                ])
                
            t = Table(data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))
            elements.append(t)
            
            doc.build(elements)
            return response

        return Response({'error': 'Invalid type'}, status=400)
